from io import BytesIO

from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file. Returns empty string on error."""
    if not file_bytes:
        return ""
    try:
        doc = Document(BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""
